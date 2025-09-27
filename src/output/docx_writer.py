#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de documents DOCX professionnels pour les bilans psychomoteurs
"""

import asyncio
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Cm, Inches, Pt, RGBColor

from config.settings import settings

# Import seulement les classes qui existent vraiment
from src.output.formatters import DocumentMetadata, StructureAnalyzer, TextFormatter
from src.utils.logging import get_logger


class DocumentTheme:
    """Thème visuel pour les documents"""

    def __init__(self, name: str = "professional"):
        self.name = name

        if name == "professional":
            self.primary_color = RGBColor(31, 78, 121)  # Bleu professionnel
            self.secondary_color = RGBColor(68, 114, 155)  # Bleu clair
            self.accent_color = RGBColor(91, 155, 213)  # Bleu accent
            self.text_color = RGBColor(51, 51, 51)  # Gris foncé
            self.light_gray = RGBColor(242, 242, 242)  # Gris très clair

        elif name == "medical":
            self.primary_color = RGBColor(22, 87, 67)  # Vert médical
            self.secondary_color = RGBColor(46, 125, 99)  # Vert clair
            self.accent_color = RGBColor(82, 190, 128)  # Vert accent
            self.text_color = RGBColor(33, 37, 41)  # Noir doux
            self.light_gray = RGBColor(248, 249, 250)  # Gris très clair

        elif name == "pediatric":
            self.primary_color = RGBColor(88, 86, 186)  # Violet doux
            self.secondary_color = RGBColor(126, 124, 209)  # Violet clair
            self.accent_color = RGBColor(165, 163, 232)  # Violet accent
            self.text_color = RGBColor(52, 58, 64)  # Gris foncé
            self.light_gray = RGBColor(244, 245, 247)  # Gris très clair


class SimpleQualityReporter:
    """Version simplifiée du générateur de rapports qualité"""

    def __init__(self):
        self.logger = get_logger(self.__class__.__name__)

    async def generate_report(
        self, sections_text: Dict[str, str], metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Génère un rapport qualité basique"""

        total_words = sum(len(text.split()) for text in sections_text.values())
        quality_score = metadata.get("quality_score", 0.0)

        report = {
            "Résumé": f"Bilan de {total_words} mots avec score qualité de {quality_score:.1%}",
            "Métriques": {
                "Nombre de mots": total_words,
                "Sections générées": len(sections_text),
                "Score qualité": f"{quality_score:.1%}",
                "Modèle utilisé": metadata.get("model_name", "N/A"),
            },
            "Sections": {
                name: f"{len(content.split())} mots"
                for name, content in sections_text.items()
            },
        }

        return report


class EnhancedDocxWriter:
    """
    Générateur de documents DOCX avancé pour bilans psychomoteurs

    Fonctionnalités :
    - Mise en page professionnelle avec thèmes
    - Structure hiérarchique automatique
    - Formatage intelligent du contenu
    - En-têtes et pieds de page personnalisés
    - Table des matières automatique
    - Métadonnées et propriétés du document
    - Export avec rapport qualité optionnel
    """

    def __init__(self, theme: str = "professional"):
        self.theme = DocumentTheme(theme)
        self.formatter = TextFormatter()
        self.structure_analyzer = StructureAnalyzer()
        self.quality_reporter = (
            SimpleQualityReporter()
        )  # Utiliser la version simplifiée
        self.logger = get_logger(self.__class__.__name__)

        # Configuration par défaut
        self.page_margins = {
            "top": Cm(2.5),
            "bottom": Cm(2.0),
            "left": Cm(2.5),
            "right": Cm(2.0),
        }

        self.font_config = {"main": "Calibri", "heading": "Calibri", "mono": "Consolas"}

        self.logger.info(f"DocxWriter initialisé avec thème: {theme}")

    async def build_document(
        self,
        sections_text: Dict[str, str],
        output_path: Path,
        title: str = "Bilan Psychomoteur",
        metadata: Optional[Dict[str, Any]] = None,
        include_toc: bool = True,
        include_quality_report: bool = None,
    ) -> None:
        """
        Génère un document DOCX complet

        Args:
            sections_text: Dictionnaire des sections avec leur contenu
            output_path: Chemin de sortie du document
            title: Titre du document
            metadata: Métadonnées additionnelles
            include_toc: Inclure une table des matières
            include_quality_report: Inclure un rapport qualité
        """
        self.logger.info(f"Génération document: {output_path}")

        # Configuration qualité par défaut
        if include_quality_report is None:
            include_quality_report = settings.output.include_quality_report

        # Analyse de la structure
        document_structure = self.structure_analyzer.analyze_sections(sections_text)

        # Création du document
        doc = Document()

        # Configuration de base
        self._setup_document_properties(doc, title, metadata)
        self._setup_styles(doc)
        self._setup_page_layout(doc)

        # Construction du contenu
        await self._build_header(doc, title, metadata)

        if include_toc:
            self._add_table_of_contents(doc, document_structure)

        await self._build_main_content(doc, sections_text, document_structure)

        if include_quality_report and metadata:
            await self._add_quality_report(doc, sections_text, metadata)

        self._add_footer(doc, metadata)

        # Sauvegarde
        doc.save(str(output_path))

        self.logger.info(f"Document généré: {output_path}")

    def _setup_document_properties(
        self, doc: Document, title: str, metadata: Optional[Dict[str, Any]]
    ) -> None:
        """Configure les propriétés du document"""
        core_props = doc.core_properties

        core_props.title = title
        core_props.subject = "Bilan Psychomoteur"
        core_props.author = "Générateur PBG v2.0"
        core_props.category = "Santé - Psychomotricité"
        core_props.keywords = "psychomotricité,bilan,évaluation,santé"
        core_props.comments = "Document généré automatiquement"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()

        if metadata:
            if metadata.get("generation_id"):
                core_props.identifier = metadata["generation_id"]
            if metadata.get("model_name"):
                core_props.comments += f" - Modèle: {metadata['model_name']}"

    def _setup_styles(self, doc: Document) -> None:
        """Configure les styles du document"""
        styles = doc.styles

        # Style titre principal
        if "Title Custom" not in [s.name for s in styles]:
            title_style = styles.add_style("Title Custom", WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = self.font_config["heading"]
            title_font.size = Pt(20)
            title_font.bold = True
            title_font.color.rgb = self.theme.primary_color

            title_para = title_style.paragraph_format
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.space_after = Pt(24)
            title_para.space_before = Pt(12)

        # Style sous-titre
        if "Subtitle Custom" not in [s.name for s in styles]:
            subtitle_style = styles.add_style(
                "Subtitle Custom", WD_STYLE_TYPE.PARAGRAPH
            )
            subtitle_font = subtitle_style.font
            subtitle_font.name = self.font_config["heading"]
            subtitle_font.size = Pt(14)
            subtitle_font.italic = True
            subtitle_font.color.rgb = self.theme.secondary_color

            subtitle_para = subtitle_style.paragraph_format
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.space_after = Pt(18)

        # Style heading personnalisé
        for level in range(1, 4):
            style_name = f"Heading {level} Custom"
            if style_name not in [s.name for s in styles]:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_font = heading_style.font
                heading_font.name = self.font_config["heading"]
                heading_font.size = Pt(16 - level)
                heading_font.bold = True
                heading_font.color.rgb = self.theme.primary_color

                heading_para = heading_style.paragraph_format
                heading_para.space_before = Pt(18 - level * 2)
                heading_para.space_after = Pt(12 - level * 2)
                heading_para.keep_with_next = True

        # Style corps de texte
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = self.font_config["main"]
        normal_font.size = Pt(11)
        normal_font.color.rgb = self.theme.text_color

        normal_para = normal_style.paragraph_format
        normal_para.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_para.line_spacing = 1.15
        normal_para.space_after = Pt(6)

        # Style pour les listes
        if "List Custom" not in [s.name for s in styles]:
            list_style = styles.add_style("List Custom", WD_STYLE_TYPE.PARAGRAPH)
            list_font = list_style.font
            list_font.name = self.font_config["main"]
            list_font.size = Pt(11)

            list_para = list_style.paragraph_format
            list_para.left_indent = Cm(0.63)
            list_para.space_after = Pt(3)

        # Style pour les informations importantes
        if "Important" not in [s.name for s in styles]:
            important_style = styles.add_style("Important", WD_STYLE_TYPE.CHARACTER)
            important_font = important_style.font
            important_font.bold = True
            important_font.color.rgb = self.theme.accent_color

    def _setup_page_layout(self, doc: Document) -> None:
        """Configure la mise en page"""
        for section in doc.sections:
            # Marges
            section.top_margin = self.page_margins["top"]
            section.bottom_margin = self.page_margins["bottom"]
            section.left_margin = self.page_margins["left"]
            section.right_margin = self.page_margins["right"]

            # En-tête et pied de page
            section.header_distance = Cm(1.27)
            section.footer_distance = Cm(1.27)

    async def _build_header(
        self, doc: Document, title: str, metadata: Optional[Dict[str, Any]]
    ) -> None:
        """Construit l'en-tête du document"""

        # Titre principal
        title_para = doc.add_paragraph(title, style="Title Custom")

        # Sous-titre avec informations contextuelles
        subtitle_parts = []

        if metadata:
            if metadata.get("generation_date"):
                try:
                    gen_date = datetime.fromisoformat(
                        metadata["generation_date"].replace("Z", "+00:00")
                    )
                    subtitle_parts.append(
                        f"Généré le {gen_date.strftime('%d/%m/%Y à %H:%M')}"
                    )
                except:
                    subtitle_parts.append(
                        f"Généré le {datetime.now().strftime('%d/%m/%Y')}"
                    )

            if metadata.get("model_name"):
                subtitle_parts.append(f"Modèle: {metadata['model_name']}")

            if metadata.get("quality_score"):
                subtitle_parts.append(f"Score qualité: {metadata['quality_score']:.1%}")

        if subtitle_parts:
            subtitle = " • ".join(subtitle_parts)
            doc.add_paragraph(subtitle, style="Subtitle Custom")

        # Ligne de séparation
        self._add_horizontal_line(doc)

    def _add_table_of_contents(
        self, doc: Document, document_structure: Dict[str, Any]
    ) -> None:
        """Ajoute une table des matières"""

        # Titre de la table des matières
        toc_title = doc.add_paragraph("Table des matières", style="Heading 1 Custom")

        # Construction de la table
        for i, section_name in enumerate(settings.section_order, 1):
            if section_name in document_structure["sections"]:
                section_info = document_structure["sections"][section_name]

                # Entrée principale
                toc_entry = doc.add_paragraph(style="List Custom")
                toc_entry.add_run(f"{i}. {section_name}")

                # Estimation du nombre de pages (approximative)
                word_count = section_info.get("word_count", 0)
                if word_count > 200:
                    toc_entry.add_run(" " + "." * 40 + f" {word_count // 150 + 1}p")

        # Saut de page après la table des matières
        doc.add_page_break()

    async def _build_main_content(
        self,
        doc: Document,
        sections_text: Dict[str, str],
        document_structure: Dict[str, Any],
    ) -> None:
        """Construit le contenu principal du document"""

        for section_name in settings.section_order:
            if section_name not in sections_text:
                continue

            section_content = sections_text[section_name]
            if not section_content or section_content.strip() == "":
                continue

            # Titre de section
            section_heading = doc.add_paragraph(section_name, style="Heading 1 Custom")

            # Traitement du contenu selon le type de section
            if section_name == "Évaluation psychomotrice":
                await self._format_evaluation_section(doc, section_content)
            elif section_name == "Tests / outils utilisés":
                self._format_tests_section(doc, section_content)
            elif "recommandation" in section_name.lower():
                self._format_recommendations_section(doc, section_content)
            else:
                self._format_standard_section(doc, section_content)

            # Espacement entre sections
            doc.add_paragraph()

    async def _format_evaluation_section(self, doc: Document, content: str) -> None:
        """Formate spécifiquement la section d'évaluation psychomotrice"""

        # Analyse du contenu pour détecter les sous-sections
        subsections = self.formatter.detect_subsections(
            content, settings.eval_subsections
        )

        if subsections:
            # Formatage par sous-sections
            for subsection_name, subsection_content in subsections.items():
                if subsection_content.strip():
                    # Sous-titre
                    doc.add_paragraph(subsection_name, style="Heading 2 Custom")

                    # Contenu formaté
                    formatted_content = self.formatter.format_clinical_text(
                        subsection_content
                    )
                    self._add_formatted_paragraphs(doc, formatted_content)
        else:
            # Formatage standard si pas de sous-sections détectées
            formatted_content = self.formatter.format_clinical_text(content)
            self._add_formatted_paragraphs(doc, formatted_content)

    def _format_tests_section(self, doc: Document, content: str) -> None:
        """Formate la section des tests et outils"""

        # Détection de listes
        test_items = self.formatter.extract_list_items(content)

        if test_items:
            # Formatage en liste à puces
            for item in test_items:
                para = doc.add_paragraph(style="List Custom")
                para.add_run("• ")
                para.add_run(item.strip())
        else:
            # Formatage standard
            formatted_content = self.formatter.format_clinical_text(content)
            self._add_formatted_paragraphs(doc, formatted_content)

    def _format_recommendations_section(self, doc: Document, content: str) -> None:
        """Formate les sections de recommandations"""

        # Mise en évidence des recommandations importantes
        formatted_content = self.formatter.highlight_recommendations(content)

        paragraphs = formatted_content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()

                # Détection des mots-clés importants
                if any(
                    keyword in para_text.lower()
                    for keyword in ["recommandé", "conseillé", "urgent", "priorité"]
                ):
                    # Mise en évidence
                    para.add_run(para_text.strip()).font.bold = True
                else:
                    para.add_run(para_text.strip())

    def _format_standard_section(self, doc: Document, content: str) -> None:
        """Formatage standard pour les autres sections"""
        formatted_content = self.formatter.format_clinical_text(content)
        self._add_formatted_paragraphs(doc, formatted_content)

    def _add_formatted_paragraphs(self, doc: Document, content: str) -> None:
        """Ajoute des paragraphes formatés au document"""

        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph()

                # Application du formatage intelligent
                formatted_runs = self.formatter.parse_emphasis(para_text)

                for text, emphasis in formatted_runs:
                    run = para.add_run(text)

                    if "bold" in emphasis:
                        run.font.bold = True
                    if "italic" in emphasis:
                        run.font.italic = True
                    if "important" in emphasis:
                        run.style = doc.styles["Important"]

    async def _add_quality_report(
        self, doc: Document, sections_text: Dict[str, str], metadata: Dict[str, Any]
    ) -> None:
        """Ajoute un rapport qualité en annexe"""

        # Saut de page pour l'annexe
        doc.add_page_break()

        # Titre de l'annexe
        doc.add_paragraph("Annexe - Rapport Qualité", style="Heading 1 Custom")

        # Génération du rapport
        quality_report = await self.quality_reporter.generate_report(
            sections_text, metadata
        )

        # Ajout du contenu du rapport
        for section_title, section_content in quality_report.items():
            doc.add_paragraph(section_title, style="Heading 2 Custom")

            if isinstance(section_content, dict):
                # Formatage de métriques
                for key, value in section_content.items():
                    para = doc.add_paragraph(style="List Custom")
                    para.add_run(f"• {key}: ")
                    para.add_run(str(value)).font.bold = True
            else:
                # Texte simple
                doc.add_paragraph(str(section_content))

    def _add_footer(self, doc: Document, metadata: Optional[Dict[str, Any]]) -> None:
        """Ajoute un pied de page"""

        section = doc.sections[0]
        footer = section.footer

        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contenu du pied de page
        footer_text = "Document confidentiel - Usage professionnel uniquement"

        if metadata and metadata.get("generation_id"):
            footer_text += f" • ID: {metadata['generation_id']}"

        footer_para.add_run(footer_text).font.size = Pt(9)
        footer_para.add_run().font.color.rgb = self.theme.secondary_color

    def _add_horizontal_line(self, doc: Document) -> None:
        """Ajoute une ligne horizontale décorative"""

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Création d'une ligne via un tableau minimal
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        cell.text = ""

        # Formatage de la ligne
        for border_name in ["top", "bottom", "left", "right"]:
            border = getattr(cell._element.tcPr.tcBorders, border_name)
            if border_name in ["top", "bottom"]:
                border.val = "single"
                border.sz = 4
                border.color = self.theme.primary_color.rgb

    def export_with_options(
        self, sections_text: Dict[str, str], output_path: Path, options: Dict[str, Any]
    ) -> DocumentMetadata:
        """
        Export avancé avec options personnalisées

        Args:
            sections_text: Contenu des sections
            output_path: Chemin de sortie
            options: Options d'export

        Returns:
            Métadonnées du document généré
        """

        # Options par défaut
        default_options = {
            "title": "Bilan Psychomoteur",
            "theme": "professional",
            "include_toc": True,
            "include_quality_report": True,
            "page_numbers": True,
            "watermark": None,
            "protection": None,
        }

        export_options = {**default_options, **options}

        # Application du thème
        if export_options["theme"] != self.theme.name:
            self.theme = DocumentTheme(export_options["theme"])

        # Génération avec options
        asyncio.run(
            self.build_document(
                sections_text=sections_text,
                output_path=output_path,
                title=export_options["title"],
                include_toc=export_options["include_toc"],
                include_quality_report=export_options["include_quality_report"],
            )
        )

        # Retour des métadonnées
        return DocumentMetadata(
            file_path=output_path,
            title=export_options["title"],
            theme=export_options["theme"],
            creation_date=datetime.now(),
            word_count=sum(len(text.split()) for text in sections_text.values()),
            section_count=len(sections_text),
        )
